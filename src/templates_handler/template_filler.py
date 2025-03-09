import re
from docx import Document

def update_docx_template(docx_path, output_path, extracted_data):
    """
    Abre la plantilla DOCX, busca las llaves {} en el texto y las reemplaza con los valores correspondientes de extracted_data.
    """
    doc = Document(docx_path)
    
    # Mapeo de las llaves en la plantilla a los datos extraídos
    key_mapping = {
        "matin": "matricula",
        "cedcas": "cedula_catastral",
        "ubipre": "ubicacion_predio",
        "dirinm": "direccion_inmueble",
        "descinm": "descripcion_inmueble",
        "val_letras": "valorVenta_letras",
        "valinm": "valorVenta",
        "nomvende": "nombre_vendedor",
        "numced_vend": "num_doc_vendedor",
        "exped_vend": "lugar_expe_vendedor",
        "estcivil_vend": "estadoCivil_vendedor",
        "conosin_vend": "sociedad_vendedor",
        "nomcompra": "nombre_comprador",
        "numced_compra": "num_doc_comprador",
        "exped_compra": "lugar_expe_comprador",
        "estcivil_compra": "estadoCivil_comprador",
        "conosin_compra": "sociedad_comprador",
        "lindpredio": "linderos_predio"
    }
    
    def format_value(value):
        """Convierte cualquier tipo de dato en una cadena correctamente formateada."""
        if isinstance(value, set):
            return " ".join(map(str, value)) if value else ""
        if isinstance(value, list):
            return ", ".join(map(str, value))
        if isinstance(value, (int, float)):
            return str(value)
        if value is None:
            return ""
        return str(value)
    
    # Recorre cada párrafo del documento y reemplaza los valores
    for paragraph in doc.paragraphs:
        for key, mapped_key in key_mapping.items():
            placeholder = f"{{{key}}}"  # Crea el marcador a buscar en la plantilla
            if placeholder in paragraph.text:
                value = format_value(extracted_data.get(mapped_key, "No especificado"))
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Recorre las tablas en el documento y reemplaza los valores en celdas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, mapped_key in key_mapping.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in cell.text:
                        value = format_value(extracted_data.get(mapped_key, "No especificado"))
                        cell.text = cell.text.replace(placeholder, value)
    
    # Guarda el documento con los datos reemplazados
    doc.save(output_path)
    print(f"Documento guardado en {output_path}")
